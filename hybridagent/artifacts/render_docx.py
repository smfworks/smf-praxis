"""Optional DOCX renderer for professional artifacts."""
from __future__ import annotations

import io
from datetime import datetime, timezone
from typing import Any

from hybridagent.artifacts.models import (
    ArtifactDocument,
    FigureBlock,
    ListBlock,
    PageBreakBlock,
    ParagraphBlock,
    Section,
    TableBlock,
)
from hybridagent.artifacts.render_common import (
    ArtifactRenderError,
    checked_assets,
    image_kind,
    normalize_zip_package,
    require_backend,
)

_FIXED_TIME = datetime(2000, 1, 1, tzinfo=timezone.utc)


def _citations(values: tuple[str, ...]) -> str:
    return "" if not values else " [" + ", ".join(values) + "]"


def _add_section(doc: Any, section: Section, assets: dict[str, bytes]) -> None:
    from docx.oxml import OxmlElement  # type: ignore[import-untyped]
    from docx.oxml.ns import qn  # type: ignore[import-untyped]

    doc.add_heading(section.title, level=section.level)
    for block in section.blocks:
        if type(block) is ParagraphBlock:
            doc.add_paragraph(block.text + _citations(block.citation_ids))
        elif type(block) is ListBlock:
            style = "List Number" if block.ordered else "List Bullet"
            for index, item in enumerate(block.items):
                suffix = _citations(block.citation_ids) if index == len(block.items) - 1 else ""
                doc.add_paragraph(item + suffix, style=style)
        elif type(block) is TableBlock:
            if block.caption:
                doc.add_paragraph(block.caption + _citations(block.citation_ids), style="Caption")
            table = doc.add_table(rows=1, cols=len(block.columns))
            table.style = "Table Grid"
            for index, value in enumerate(block.columns):
                table.rows[0].cells[index].text = value
            # Repeat the semantic header row in supporting readers.
            header = OxmlElement("w:tblHeader")
            header.set(qn("w:val"), "true")
            table.rows[0]._tr.get_or_add_trPr().append(header)
            for row in block.rows:
                cells = table.add_row().cells
                for index, value in enumerate(row):
                    cells[index].text = value
        elif type(block) is FigureBlock:
            payload = assets[block.asset_id]
            if image_kind(payload) == "svg":
                raise ArtifactRenderError("DOCX rendering requires PNG or JPEG figure assets")
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            picture = run.add_picture(io.BytesIO(payload))
            picture._inline.docPr.set("descr", block.alt_text)
            picture._inline.docPr.set("title", block.caption or block.alt_text)
            doc.add_paragraph((block.caption or block.alt_text) + _citations(block.citation_ids), style="Caption")
        elif type(block) is PageBreakBlock:
            doc.add_page_break()


def render_docx(
    document: ArtifactDocument, assets: dict[str, bytes] | None = None
) -> bytes:
    require_backend("docx", "python-docx")
    from docx import Document  # type: ignore[import-untyped]
    from docx.enum.section import WD_SECTION  # type: ignore[import-untyped]
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import-untyped]

    resolved = checked_assets(document, assets)
    doc = Document()
    properties = doc.core_properties
    properties.title = document.metadata.title
    properties.subject = document.metadata.subject
    properties.author = document.metadata.created_by
    properties.last_modified_by = "Praxis Artifact Studio"
    properties.created = _FIXED_TIME
    properties.modified = _FIXED_TIME
    properties.category = document.metadata.document_type
    properties.comments = f"artifact={document.artifact_id}; sha256={document.content_hash()}"
    title = doc.add_heading(document.metadata.title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Classification: {document.metadata.confidentiality.upper()}")
    if document.metadata.subject:
        doc.add_paragraph(f"Subject: {document.metadata.subject}")
    for section in document.sections:
        _add_section(doc, section, resolved)
    if document.appendices:
        doc.add_section(WD_SECTION.NEW_PAGE)
        doc.add_heading("Appendices", 1)
        for appendix in document.appendices:
            _add_section(doc, appendix, resolved)
    if document.sources:
        doc.add_heading("Source Manifest", 1)
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        headers = ("Source", "Version", "SHA-256", "Title", "URI")
        for index, value in enumerate(headers):
            table.rows[0].cells[index].text = value
        for source in document.sources:
            cells = table.add_row().cells
            for index, value in enumerate((source.source_id, source.source_version_id,
                                            source.content_hash, source.title,
                                            source.canonical_uri)):
                cells[index].text = value
    if document.revisions:
        doc.add_heading("Revision History", 1)
        for revision in document.revisions:
            doc.add_paragraph(
                f"{revision.sequence}. {revision.summary} — {revision.author_id}, {revision.created_at}"
            )
    if document.reviews:
        doc.add_heading("Professional Reviews", 1)
        for review in document.reviews:
            doc.add_paragraph(
                f"{review.review_type}: {review.decision} — {review.reviewer_id}, {review.reviewed_at}"
            )
    if document.signatures:
        doc.add_heading("Signatures", 1)
        for signature in document.signatures:
            doc.add_paragraph(
                f"{signature.signer_id} ({signature.role}) — {signature.meaning}, {signature.signed_at}"
            )
    for doc_section in doc.sections:
        footer = doc_section.footer.paragraphs[0]
        footer.text = f"{document.metadata.confidentiality.upper()} — {document.artifact_id}"
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    output = io.BytesIO()
    doc.save(output)
    return normalize_zip_package(output.getvalue())
